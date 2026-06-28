#!/usr/bin/env python3
"""
Hernoem WK-pouleformulieren op basis van de naam in het formulier.

Standaard wordt alleen een voorbeeld getoond.
Gebruik --uitvoeren om de bestanden echt te hernoemen.

Voorbeeld:
    py -3 hernoem_formulieren.py
    py -3 hernoem_formulieren.py --uitvoeren

Standaardmap:
    data/formulieren/
"""

from __future__ import annotations

import argparse
import logging
import re
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
DEFAULT_FOLDER = ROOT / "data" / "formulieren"

logging.getLogger("pypdf").setLevel(logging.ERROR)


def schone_bestandsnaam(naam: str) -> str:
    """Maak een Windows-veilige bestandsnaam."""
    naam = re.sub(r'[<>:"/\\|?*]', "", naam)
    naam = re.sub(r"\s+", " ", naam).strip()
    naam = naam.rstrip(". ")

    if not naam:
        raise ValueError("lege naam na opschonen")

    # Vermijd gereserveerde Windows-bestandsnamen.
    verboden = {
        "CON", "PRN", "AUX", "NUL",
        *(f"COM{i}" for i in range(1, 10)),
        *(f"LPT{i}" for i in range(1, 10)),
    }
    if naam.upper() in verboden:
        naam = f"Deelnemer {naam}"

    return naam


def lees_naam_docx(pad: Path) -> str:
    document = Document(pad)

    # Zoek eerst gericht naar een cel met "Naam".
    for table in document.tables:
        for row in table.rows:
            cells = [
                " ".join(cell.text.replace("\n", " ").split()).strip()
                for cell in row.cells
            ]

            for index, waarde in enumerate(cells):
                if waarde.casefold().rstrip(":") == "naam":
                    if index + 1 < len(cells) and cells[index + 1]:
                        return cells[index + 1]

    # Fallback: zoek in gewone tekst.
    volledige_tekst = "\n".join(
        paragraph.text for paragraph in document.paragraphs
    )
    match = re.search(
        r"\bNaam\s*:\s*(.+?)(?:\n|Voor de poulefase|$)",
        volledige_tekst,
        flags=re.I,
    )
    if match:
        return match.group(1).strip()

    raise ValueError("naam niet gevonden in DOCX")


def lees_naam_pdf(pad: Path) -> str:
    reader = PdfReader(pad)
    tekst = "\n".join(page.extract_text() or "" for page in reader.pages)
    compacte_tekst = " ".join(tekst.split())

    match = re.search(
        r"\bNaam\s*:\s*(.*?)\s+Voor de poulefase",
        compacte_tekst,
        flags=re.I,
    )
    if match:
        return match.group(1).strip()

    match = re.search(
        r"\bNaam\s*:\s*([A-Za-zÀ-ÿ0-9 .'\-/&()]+)",
        compacte_tekst,
        flags=re.I,
    )
    if match:
        return match.group(1).strip()

    raise ValueError("naam niet gevonden in PDF")


def lees_naam(pad: Path) -> str:
    extensie = pad.suffix.casefold()

    if extensie == ".docx":
        return lees_naam_docx(pad)
    if extensie == ".pdf":
        return lees_naam_pdf(pad)

    raise ValueError(f"niet-ondersteund bestandstype: {pad.suffix}")


def unieke_bestemming(
    map_pad: Path,
    gewenste_naam: str,
    extensie: str,
    bron: Path,
) -> Path:
    kandidaat = map_pad / f"{gewenste_naam}{extensie}"

    # Als het al exact hetzelfde bestand is, niets aanpassen.
    if kandidaat.resolve() == bron.resolve():
        return kandidaat

    teller = 2
    while kandidaat.exists():
        kandidaat = map_pad / f"{gewenste_naam} ({teller}){extensie}"
        teller += 1

    return kandidaat


def vind_bestanden(map_pad: Path) -> list[Path]:
    return sorted(
        pad
        for pad in map_pad.rglob("*")
        if pad.is_file()
        and pad.suffix.casefold() in {".docx", ".pdf"}
        and not pad.name.startswith("~$")
    )


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--map",
        type=Path,
        default=DEFAULT_FOLDER,
        help="map met formulieren",
    )
    parser.add_argument(
        "--uitvoeren",
        action="store_true",
        help="hernoem bestanden echt; zonder deze optie is het alleen een voorbeeld",
    )
    parser.add_argument(
        "--zonder-prefix",
        action="store_true",
        help="gebruik alleen de naam, zonder 'Formulier-A - ' ervoor",
    )
    args = parser.parse_args(argv)

    if not args.map.exists():
        print(f"Map niet gevonden: {args.map}", file=sys.stderr)
        return 1

    bestanden = vind_bestanden(args.map)
    if not bestanden:
        print(f"Geen DOCX- of PDF-bestanden gevonden in: {args.map}")
        return 1

    gewijzigd = 0
    fouten = 0

    for bron in bestanden:
        try:
            naam = schone_bestandsnaam(lees_naam(bron))
            basisnaam = naam if args.zonder_prefix else f"Formulier-A - {naam}"

            bestemming = unieke_bestemming(
                map_pad=bron.parent,
                gewenste_naam=basisnaam,
                extensie=bron.suffix.lower(),
                bron=bron,
            )

            if bestemming.resolve() == bron.resolve():
                print(f"OK   {bron.name} is al goed")
                continue

            if args.uitvoeren:
                bron.rename(bestemming)
                print(f"HERNOEMD  {bron.name} -> {bestemming.name}")
            else:
                print(f"VOORBEELD {bron.name} -> {bestemming.name}")

            gewijzigd += 1

        except Exception as fout:
            fouten += 1
            print(f"FOUT {bron.name}: {fout}", file=sys.stderr)

    print()
    if args.uitvoeren:
        print(f"Klaar: {gewijzigd} bestand(en) hernoemd.")
    else:
        print(f"Voorbeeld klaar: {gewijzigd} bestand(en) zouden worden hernoemd.")
        print("Voer opnieuw uit met --uitvoeren om dit echt toe te passen.")

    if fouten:
        print(f"Let op: {fouten} bestand(en) konden niet worden verwerkt.")

    return 0 if gewijzigd or not fouten else 1


if __name__ == "__main__":
    raise SystemExit(main())
