#!/usr/bin/env python3
"""Hernoem fase-2-formulieren naar een vaste bestandsnaam.

Voorbeelden:

    Formulier B zestiende finales (3).docx
    -> Formulier fase2 - Sander Stok.docx

    Formulier B zestiende finales V2.pdf
    -> Formulier fase2 - Ravi Biekram.pdf

Het script leest de naam uit ieder DOCX- of PDF-formulier en koppelt deze aan
de officiële naam in data/fase2/fase1_punten.csv.

Gebruik vanuit de repository-root:

    python hernoem_fase2_formulieren.py

Standaard worden bestanden echt hernoemd. Gebruik --dry-run om eerst alleen
te bekijken wat er zou veranderen:

    python hernoem_fase2_formulieren.py --dry-run
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def normaliseer(waarde: object) -> str:
    tekst = unicodedata.normalize("NFD", str(waarde or ""))
    tekst = "".join(
        teken
        for teken in tekst
        if unicodedata.category(teken) != "Mn"
    )
    tekst = tekst.casefold()
    tekst = re.sub(r"[^a-z0-9]+", " ", tekst)
    return " ".join(tekst.split())


def veilige_bestandsnaam(naam: str) -> str:
    naam = re.sub(r'[<>:"/\\|?*]+', "", naam)
    naam = re.sub(r"\s+", " ", naam).strip()
    return naam.rstrip(". ")


def lees_csv(pad: Path) -> list[dict[str, str]]:
    if not pad.exists():
        raise FileNotFoundError(f"Bestand niet gevonden: {pad}")

    with pad.open("r", encoding="utf-8-sig", newline="") as bestand:
        return list(csv.DictReader(bestand))


def lees_naam_docx(pad: Path) -> str:
    document = Document(pad)

    for tabel in document.tables:
        for rij in tabel.rows:
            cellen = [cel.text.strip() for cel in rij.cells]

            if (
                len(cellen) >= 2
                and normaliseer(cellen[0]).startswith("naam")
            ):
                return cellen[1].strip()

    for alinea in document.paragraphs:
        match = re.search(
            r"\bnaam\s*:\s*(.+)$",
            alinea.text,
            flags=re.I,
        )
        if match:
            return match.group(1).strip()

    return ""


def lees_naam_pdf(pad: Path) -> str:
    reader = PdfReader(str(pad))
    tekst = "\n".join(
        pagina.extract_text() or ""
        for pagina in reader.pages
    )

    match = re.search(
        r"\bnaam\s*:\s*([^\n\r]+)",
        tekst,
        flags=re.I,
    )

    return match.group(1).strip() if match else ""


def lees_formuliernaam(pad: Path) -> str:
    suffix = pad.suffix.casefold()

    if suffix == ".docx":
        return lees_naam_docx(pad)

    if suffix == ".pdf":
        return lees_naam_pdf(pad)

    return ""


def match_officiele_naam(
    formuliernaam: str,
    bekende_namen: dict[str, str],
) -> str:
    sleutel = normaliseer(formuliernaam)

    if sleutel in bekende_namen:
        return bekende_namen[sleutel]

    # Alleen een voornaam ingevuld, bijvoorbeeld Sander -> Sander Stok.
    delen = sleutel.split()
    if len(delen) == 1:
        kandidaten = [
            echte_naam
            for genormaliseerde_naam, echte_naam
            in bekende_namen.items()
            if genormaliseerde_naam.split()[:1] == delen
        ]

        if len(kandidaten) == 1:
            return kandidaten[0]

    raise ValueError(
        f"Naam {formuliernaam!r} kan niet uniek worden gekoppeld "
        "aan fase1_punten.csv."
    )


def uniek_doelpad(
    map_pad: Path,
    basisnaam: str,
    suffix: str,
    bron: Path,
) -> Path:
    kandidaat = map_pad / f"{basisnaam}{suffix}"

    if kandidaat == bron:
        return kandidaat

    if not kandidaat.exists():
        return kandidaat

    teller = 2
    while True:
        kandidaat = map_pad / f"{basisnaam} ({teller}){suffix}"
        if not kandidaat.exists() or kandidaat == bron:
            return kandidaat
        teller += 1


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--data-dir",
        type=Path,
        default=Path("data/fase2"),
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Toon wijzigingen zonder bestanden te hernoemen.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    data_dir: Path = args.data_dir
    formulieren_dir = data_dir / "formulieren"
    fase1_pad = data_dir / "fase1_punten.csv"

    try:
        fase1_rijen = lees_csv(fase1_pad)
        bekende_namen: dict[str, str] = {}

        for rij in fase1_rijen:
            naam = str(rij.get("naam", "")).strip()
            if naam:
                bekende_namen[normaliseer(naam)] = naam

        if not bekende_namen:
            raise ValueError(
                "fase1_punten.csv bevat geen geldige namen."
            )

        if not formulieren_dir.exists():
            raise FileNotFoundError(
                f"Formulierenmap niet gevonden: {formulieren_dir}"
            )

        bestanden = sorted(
            [
                pad
                for pad in formulieren_dir.iterdir()
                if pad.is_file()
                and pad.suffix.casefold() in {".docx", ".pdf"}
            ],
            key=lambda pad: pad.name.casefold(),
        )

        if not bestanden:
            print("Geen DOCX- of PDF-formulieren gevonden.")
            return 0

        hernoemd = 0
        overgeslagen = 0
        fouten = 0

        for bron in bestanden:
            try:
                formuliernaam = lees_formuliernaam(bron)

                if not formuliernaam:
                    raise ValueError(
                        "geen naam in het formulier gevonden"
                    )

                officiele_naam = match_officiele_naam(
                    formuliernaam,
                    bekende_namen,
                )

                veilige_naam = veilige_bestandsnaam(officiele_naam)
                basisnaam = f"Formulier fase2 - {veilige_naam}"

                doel = uniek_doelpad(
                    formulieren_dir,
                    basisnaam,
                    bron.suffix.casefold(),
                    bron,
                )

                if doel == bron:
                    print(f"OK: {bron.name} is al goed.")
                    overgeslagen += 1
                    continue

                print(f"{bron.name} -> {doel.name}")

                if not args.dry_run:
                    bron.rename(doel)

                hernoemd += 1

            except Exception as fout:
                fouten += 1
                print(
                    f"FOUT: {bron.name}: {fout}",
                    file=sys.stderr,
                )

        print()
        print(f"Hernoemd: {hernoemd}")
        print(f"Al goed: {overgeslagen}")
        print(f"Fouten: {fouten}")

        if args.dry_run:
            print("Dry-run: er zijn geen bestanden aangepast.")

        return 1 if fouten else 0

    except Exception as fout:
        print(f"Fout: {fout}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
