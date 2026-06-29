#!/usr/bin/env python3
"""Importeer fase-2-formulieren uit DOCX en PDF.

Ondersteunt onder andere:
- losse scorekolommen: Nederland | 2 | Marokko | 1;
- één gecombineerde scorekolom: Nederland | Marokko | 2-1;
- PDF-tabellen over meerdere pagina's;
- lege voorspellingen;
- unieke voornaammatching;
- gedeeltelijke updates zonder andere deelnemers te verwijderen.

Benodigd:
    pip install python-docx pdfplumber
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
import unicodedata
from pathlib import Path

import pdfplumber
from docx import Document


AANTAL_WEDSTRIJDEN = 16
UITVOERVELDEN = [
    "naam",
    "wedstrijd",
    "thuis",
    "voorspeld_thuis",
    "uit",
    "voorspeld_uit",
    "winnaar_na_penalties",
    "bronbestand",
]


def normaliseer(waarde: object) -> str:
    tekst = unicodedata.normalize("NFD", str(waarde or ""))
    tekst = "".join(
        teken for teken in tekst
        if unicodedata.category(teken) != "Mn"
    )
    tekst = tekst.casefold()
    tekst = re.sub(r"[^a-z0-9]+", " ", tekst)
    return " ".join(tekst.split())


def lees_csv(pad: Path) -> list[dict[str, str]]:
    if not pad.exists():
        raise FileNotFoundError(f"Bestand niet gevonden: {pad}")

    with pad.open("r", encoding="utf-8-sig", newline="") as bestand:
        return list(csv.DictReader(bestand))


def schrijf_csv(pad: Path, rijen: list[dict[str, object]]) -> None:
    pad.parent.mkdir(parents=True, exist_ok=True)

    with pad.open("w", encoding="utf-8-sig", newline="") as bestand:
        schrijver = csv.DictWriter(
            bestand,
            fieldnames=UITVOERVELDEN,
        )
        schrijver.writeheader()
        schrijver.writerows(rijen)


def parse_scorewaarde(waarde: object) -> int | str:
    tekst = str(waarde or "").strip()

    if tekst == "":
        return ""

    if re.fullmatch(r"\d+", tekst):
        return int(tekst)

    raise ValueError(f"ongeldige scorewaarde {tekst!r}")


def parse_gecombineerde_score(waarde: object) -> tuple[int | str, int | str] | None:
    tekst = str(waarde or "").strip()

    if tekst == "":
        return None

    match = re.fullmatch(r"(\d+)\s*[-–—:]\s*(\d+)", tekst)
    if not match:
        return None

    return int(match.group(1)), int(match.group(2))


def schoon(cellen: list[object]) -> list[str]:
    return [
        re.sub(r"\s+", " ", str(cel or "")).strip()
        for cel in cellen
    ]


def laad_schema(data_dir: Path) -> dict[int, dict[str, str]]:
    rijen = lees_csv(data_dir / "uitslagen_fase2.csv")
    schema: dict[int, dict[str, str]] = {}

    for rij in rijen:
        nummertekst = str(rij.get("wedstrijd", "")).strip()
        thuis = str(rij.get("thuis", "")).strip()
        uit = str(rij.get("uit", "")).strip()

        if not nummertekst:
            continue

        nummer = int(nummertekst)
        schema[nummer] = {
            "thuis": thuis,
            "uit": uit,
        }

    ontbrekend = [
        nummer for nummer in range(1, AANTAL_WEDSTRIJDEN + 1)
        if nummer not in schema
    ]
    if ontbrekend:
        raise ValueError(
            "uitslagen_fase2.csv mist wedstrijd(en): "
            + ", ".join(map(str, ontbrekend))
        )

    return schema


def vind_naam_docx(bron: Path) -> str:
    document = Document(bron)

    for tabel in document.tables:
        for rij in tabel.rows:
            cellen = schoon([cel.text for cel in rij.cells])
            if len(cellen) >= 2 and normaliseer(cellen[0]).startswith("naam"):
                return cellen[1]

    for alinea in document.paragraphs:
        match = re.search(
            r"\bnaam\s*:\s*(.+)$",
            alinea.text,
            flags=re.I,
        )
        if match:
            return match.group(1).strip()

    return ""


def vind_naam_pdf(bron: Path) -> str:
    with pdfplumber.open(bron) as pdf:
        # Probeer eerst de echte tabelkolom te lezen.
        for pagina in pdf.pages:
            for tabel in pagina.extract_tables() or []:
                for rij in tabel or []:
                    cellen = schoon(rij or [])
                    if (
                        len(cellen) >= 2
                        and normaliseer(cellen[0]).startswith("naam")
                    ):
                        return cellen[1]

        tekst = "\n".join(
            pagina.extract_text() or ""
            for pagina in pdf.pages
        )

    match = re.search(
        r"\bnaam\s*:\s*([^\n\r]+)",
        tekst,
        flags=re.I,
    )
    return match.group(1).strip() if match else ""


def koppel_naam(
    formuliernaam: str,
    bekende_namen: dict[str, str],
) -> str:
    sleutel = normaliseer(formuliernaam)

    if sleutel in bekende_namen:
        return bekende_namen[sleutel]

    # Voornaam als unieke match, bijvoorbeeld Sander -> Sander Stok.
    delen = sleutel.split()
    if len(delen) == 1:
        kandidaten = [
            echte_naam
            for genormaliseerd, echte_naam in bekende_namen.items()
            if genormaliseerd.split()[:1] == delen
        ]
        if len(kandidaten) == 1:
            return kandidaten[0]

    # Alias/prefix, bijvoorbeeld Bolle(n)boos ->
    # Bolle(n)boos / Joost Valk.
    prefix_kandidaten = [
        echte_naam
        for genormaliseerd, echte_naam in bekende_namen.items()
        if genormaliseerd.startswith(sleutel + " ")
        or sleutel.startswith(genormaliseerd + " ")
    ]
    if len(prefix_kandidaten) == 1:
        return prefix_kandidaten[0]

    raise ValueError(
        f"naam {formuliernaam!r} kan niet uniek worden gekoppeld "
        "aan fase1_punten.csv"
    )


def parse_wedstrijdregel(
    cellen: list[object],
    schema: dict[int, dict[str, str]],
    bron: Path,
    infos: list[str],
) -> dict[str, object] | None:
    waarden = schoon(cellen)

    if not waarden or not re.fullmatch(r"\d{1,2}", waarden[0]):
        return None

    nummer = int(waarden[0])
    if nummer not in schema:
        return None

    thuis = schema[nummer]["thuis"]
    uit = schema[nummer]["uit"]

    # Vorm A: nummer | thuis | thuis-score | uit | uit-score | penalties
    # Vorm B: nummer | thuis | leeg | uit | "2-1" | penalties
    # Vorm C: nummer | thuis | uit | "2-1" | penalties
    voorspeld_thuis: int | str = ""
    voorspeld_uit: int | str = ""
    winnaar = ""

    if len(waarden) >= 5:
        gecombineerde = parse_gecombineerde_score(waarden[4])

        if gecombineerde is not None and waarden[2] == "":
            voorspeld_thuis, voorspeld_uit = gecombineerde
            winnaar = waarden[5] if len(waarden) >= 6 else ""
        else:
            try:
                voorspeld_thuis = parse_scorewaarde(waarden[2])
                voorspeld_uit = parse_scorewaarde(waarden[4])
                winnaar = waarden[5] if len(waarden) >= 6 else ""
            except ValueError:
                gecombineerde = None
                for index, waarde in enumerate(waarden[1:], start=1):
                    score = parse_gecombineerde_score(waarde)
                    if score is not None:
                        gecombineerde = score
                        winnaar = (
                            waarden[index + 1]
                            if index + 1 < len(waarden)
                            else ""
                        )
                        break

                if gecombineerde is None:
                    raise ValueError(
                        f"{bron.name}: wedstrijd {nummer} bevat geen "
                        "leesbare scorekolom"
                    )

                voorspeld_thuis, voorspeld_uit = gecombineerde
    else:
        gecombineerde = None
        for index, waarde in enumerate(waarden[1:], start=1):
            score = parse_gecombineerde_score(waarde)
            if score is not None:
                gecombineerde = score
                winnaar = (
                    waarden[index + 1]
                    if index + 1 < len(waarden)
                    else ""
                )
                break

        if gecombineerde is None:
            voorspeld_thuis = ""
            voorspeld_uit = ""
        else:
            voorspeld_thuis, voorspeld_uit = gecombineerde

    # Lege voorspellingen zijn toegestaan.
    if voorspeld_thuis == "" or voorspeld_uit == "":
        voorspeld_thuis = ""
        voorspeld_uit = ""
        winnaar = ""
        infos.append(
            f"{bron.name}: wedstrijd {nummer} heeft geen volledige score "
            "en levert 0 punten op."
        )
    elif voorspeld_thuis != voorspeld_uit:
        # Bij een niet-gelijkspel is de winnaar al uit de score af te leiden.
        winnaar = ""
    elif winnaar == "":
        infos.append(
            f"{bron.name}: wedstrijd {nummer} is gelijk voorspeld zonder "
            "penaltywinnaar; maximaal 3 punten mogelijk."
        )

    return {
        "wedstrijd": nummer,
        "thuis": thuis,
        "voorspeld_thuis": voorspeld_thuis,
        "uit": uit,
        "voorspeld_uit": voorspeld_uit,
        "winnaar_na_penalties": winnaar,
    }


def lees_docx(
    bron: Path,
    schema: dict[int, dict[str, str]],
) -> tuple[str, list[dict[str, object]], list[str]]:
    document = Document(bron)
    naam = vind_naam_docx(bron)
    infos: list[str] = []
    wedstrijden: list[dict[str, object]] = []

    for tabel in document.tables:
        for rij in tabel.rows:
            parsed = parse_wedstrijdregel(
                [cel.text for cel in rij.cells],
                schema,
                bron,
                infos,
            )
            if parsed is not None:
                wedstrijden.append(parsed)

    return naam, controleer_wedstrijden(bron, wedstrijden), infos


def lees_pdf(
    bron: Path,
    schema: dict[int, dict[str, str]],
) -> tuple[str, list[dict[str, object]], list[str]]:
    naam = vind_naam_pdf(bron)
    infos: list[str] = []
    wedstrijden: list[dict[str, object]] = []

    with pdfplumber.open(bron) as pdf:
        for pagina in pdf.pages:
            tabellen = pagina.extract_tables() or []

            for tabel in tabellen:
                for rij in tabel or []:
                    parsed = parse_wedstrijdregel(
                        rij or [],
                        schema,
                        bron,
                        infos,
                    )
                    if parsed is not None:
                        wedstrijden.append(parsed)

    # Alleen als tabelherkenning onvoldoende is, tekstregels proberen.
    gevonden = {int(item["wedstrijd"]) for item in wedstrijden}
    ontbrekend = [
        nummer
        for nummer in range(1, AANTAL_WEDSTRIJDEN + 1)
        if nummer not in gevonden
    ]

    if ontbrekend:
        with pdfplumber.open(bron) as pdf:
            tekstregels = []
            for pagina in pdf.pages:
                tekstregels.extend(
                    (pagina.extract_text() or "").splitlines()
                )

        for regel in tekstregels:
            match = re.match(r"^\s*(\d{1,2})\s+", regel)
            if not match:
                continue

            nummer = int(match.group(1))
            if nummer not in ontbrekend:
                continue

            schema_regel = schema[nummer]
            patroon = re.compile(
                rf"^\s*{nummer}\s+"
                rf"{re.escape(schema_regel['thuis'])}\s+"
                rf"(\d*)\s*"
                rf"{re.escape(schema_regel['uit'])}\s+"
                rf"(\d*)"
                rf"(?:\s+(.+))?$",
                flags=re.I,
            )
            resultaat = patroon.match(regel)
            if not resultaat:
                continue

            parsed = parse_wedstrijdregel(
                [
                    str(nummer),
                    schema_regel["thuis"],
                    resultaat.group(1) or "",
                    schema_regel["uit"],
                    resultaat.group(2) or "",
                    resultaat.group(3) or "",
                ],
                schema,
                bron,
                infos,
            )
            if parsed is not None:
                wedstrijden.append(parsed)

    return naam, controleer_wedstrijden(bron, wedstrijden), infos


def controleer_wedstrijden(
    bron: Path,
    wedstrijden: list[dict[str, object]],
) -> list[dict[str, object]]:
    per_nummer: dict[int, dict[str, object]] = {}

    for wedstrijd in wedstrijden:
        nummer = int(wedstrijd["wedstrijd"])
        per_nummer[nummer] = wedstrijd

    ontbrekend = [
        nummer
        for nummer in range(1, AANTAL_WEDSTRIJDEN + 1)
        if nummer not in per_nummer
    ]
    if ontbrekend:
        raise ValueError(
            f"{bron.name}: ontbrekende wedstrijden: "
            + ", ".join(map(str, ontbrekend))
        )

    return [
        per_nummer[nummer]
        for nummer in range(1, AANTAL_WEDSTRIJDEN + 1)
    ]


def lees_formulier(
    bron: Path,
    schema: dict[int, dict[str, str]],
) -> tuple[str, list[dict[str, object]], list[str]]:
    suffix = bron.suffix.casefold()

    if suffix == ".docx":
        return lees_docx(bron, schema)

    if suffix == ".pdf":
        return lees_pdf(bron, schema)

    raise ValueError(
        f"{bron.name}: bestandstype {bron.suffix!r} wordt niet ondersteund"
    )


def laad_bestaande_voorspellingen(
    uitvoer: Path,
) -> list[dict[str, str]]:
    if not uitvoer.exists():
        return []

    return lees_csv(uitvoer)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--data-dir",
        type=Path,
        default=Path("data/fase2"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    data_dir: Path = args.data_dir
    formulieren_dir = data_dir / "formulieren"
    fase1_pad = data_dir / "fase1_punten.csv"
    uitvoer = args.output or data_dir / "voorspellingen_fase2.csv"

    try:
        bekende_namen: dict[str, str] = {}
        for rij in lees_csv(fase1_pad):
            naam = str(rij.get("naam", "")).strip()
            if naam:
                bekende_namen[normaliseer(naam)] = naam

        schema = laad_schema(data_dir)

        formulieren = sorted(
            [
                pad for pad in formulieren_dir.iterdir()
                if pad.is_file()
                and pad.suffix.casefold() in {".docx", ".pdf"}
            ],
            key=lambda pad: (
                pad.stat().st_mtime,
                pad.name.casefold(),
            ),
        )

        bestaande_rijen = laad_bestaande_voorspellingen(uitvoer)
        per_deelnemer: dict[str, list[dict[str, object]]] = {}

        for rij in bestaande_rijen:
            naam = str(rij.get("naam", "")).strip()
            if naam:
                per_deelnemer.setdefault(
                    normaliseer(naam),
                    [],
                ).append(dict(rij))

        fouten = 0
        infos: list[str] = []

        for bron in formulieren:
            try:
                formuliernaam, wedstrijden, formulier_infos = (
                    lees_formulier(bron, schema)
                )

                if not formuliernaam:
                    raise ValueError(
                        f"{bron.name}: geen naam gevonden"
                    )

                officiele_naam = koppel_naam(
                    formuliernaam,
                    bekende_namen,
                )

                nieuwe_rijen = [
                    {
                        "naam": officiele_naam,
                        **wedstrijd,
                        "bronbestand": bron.name,
                    }
                    for wedstrijd in wedstrijden
                ]

                # Altijd precies de oude 16 regels van deze deelnemer vervangen.
                per_deelnemer[normaliseer(officiele_naam)] = nieuwe_rijen
                infos.extend(formulier_infos)

                print(
                    f"OK: {bron.name} -> {officiele_naam} "
                    f"({len(nieuwe_rijen)} wedstrijden)"
                )

            except Exception as fout:
                fouten += 1
                print(f"FOUT: {fout}", file=sys.stderr)

        alle_rijen: list[dict[str, object]] = []
        for rijen in per_deelnemer.values():
            alle_rijen.extend(rijen)

        alle_rijen.sort(
            key=lambda rij: (
                normaliseer(rij["naam"]),
                int(rij["wedstrijd"]),
            )
        )

        # Ook bij één fout worden alle succesvolle updates opgeslagen.
        schrijf_csv(uitvoer, alle_rijen)

        print(f"\nGeschreven: {uitvoer.resolve()}")
        print(
            f"Deelnemers in CSV: "
            f"{len({normaliseer(rij['naam']) for rij in alle_rijen})}"
        )

        if infos:
            print("\nINFO:")
            for info in sorted(set(infos)):
                print(f"- {info}")

        if fouten:
            print(
                f"\nLet op: {fouten} formulier(en) konden niet worden "
                "ingelezen. De andere updates zijn wel opgeslagen.",
                file=sys.stderr,
            )

        return 1 if fouten else 0

    except Exception as fout:
        print(f"Fout: {fout}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
