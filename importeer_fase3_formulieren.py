#!/usr/bin/env python3
"""Importeer fase-3-formulieren uit DOCX en PDF.

Extra gedrag:
- lege teamvelden in een volgende ronde worden automatisch ingevuld;
- teksten zoals 'Winnaar B', 'Winnaar 1' en 'Winnaar X' worden opgelost;
- de voorspelde winnaar van de vorige wedstrijd stroomt automatisch door;
- dit werkt door van achtste finales naar kwartfinales, halve finales en finale.

Bracket:
- Kwartfinale 1: winnaar B - winnaar A
- Kwartfinale 2: winnaar E - winnaar F
- Kwartfinale 3: winnaar C - winnaar D
- Kwartfinale 4: winnaar G - winnaar H
- Halve finale X: winnaar 1 - winnaar 2
- Halve finale Y: winnaar 3 - winnaar 4
- Finale F: winnaar X - winnaar Y
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
import unicodedata
from pathlib import Path

import pdfplumber
from docx import Document


FIELDS = [
    "naam",
    "ronde",
    "wedstrijd",
    "thuis",
    "voorspeld_thuis",
    "uit",
    "voorspeld_uit",
    "winnaar",
    "bronbestand",
]

ROUND_ORDER = {
    "r16": 0,
    "qf": 1,
    "sf": 2,
    "f": 3,
}

BRACKET = {
    ("qf", "1"): (("r16", "B"), ("r16", "A")),
    ("qf", "2"): (("r16", "E"), ("r16", "F")),
    ("qf", "3"): (("r16", "C"), ("r16", "D")),
    ("qf", "4"): (("r16", "G"), ("r16", "H")),
    ("sf", "X"): (("qf", "1"), ("qf", "2")),
    ("sf", "Y"): (("qf", "3"), ("qf", "4")),
    ("f", "F"): (("sf", "X"), ("sf", "Y")),
}


def normaliseer(waarde: object) -> str:
    tekst = unicodedata.normalize("NFD", str(waarde or ""))
    tekst = "".join(
        teken
        for teken in tekst
        if unicodedata.category(teken) != "Mn"
    )
    tekst = tekst.casefold()
    tekst = re.sub(r"[^a-z0-9]+", " ", tekst)
    return " ".join(tekst.split())


def schoon(waarde: object) -> str:
    return re.sub(r"\s+", " ", str(waarde or "")).strip()


def lees_csv(pad: Path) -> list[dict[str, str]]:
    if not pad.exists():
        return []

    with pad.open("r", encoding="utf-8-sig", newline="") as bestand:
        return list(csv.DictReader(bestand))


def koppel_naam(
    formuliernaam: str,
    bekende_namen: dict[str, str],
) -> str:
    sleutel = normaliseer(formuliernaam)

    if sleutel in bekende_namen:
        return bekende_namen[sleutel]

    delen = sleutel.split()

    if len(delen) == 1:
        kandidaten = [
            naam
            for genormaliseerd, naam in bekende_namen.items()
            if genormaliseerd.split()[:1] == delen
        ]

        if len(kandidaten) == 1:
            return kandidaten[0]

    kandidaten = [
        naam
        for genormaliseerd, naam in bekende_namen.items()
        if genormaliseerd.startswith(sleutel + " ")
        or sleutel.startswith(genormaliseerd + " ")
    ]

    if len(kandidaten) == 1:
        return kandidaten[0]

    raise ValueError(
        f"Naam {formuliernaam!r} kan niet uniek worden gekoppeld."
    )


def ronde_uit_tekst(tekst: str) -> str | None:
    boven = tekst.upper()

    if "ACHTSTE" in boven:
        return "r16"
    if "KWART" in boven:
        return "qf"
    if "HALVE" in boven:
        return "sf"
    if "FINALE" in boven and "HALVE" not in boven:
        return "f"

    return None


def geldige_wedstrijdcode(ronde: str, code: str) -> bool:
    if ronde == "r16":
        return code in list("ABCDEFGH")
    if ronde == "qf":
        return code in list("1234")
    if ronde == "sf":
        return code in {"X", "Y"}
    if ronde == "f":
        return code in {"", "F"}

    return False


def score_getal(waarde: str) -> int | None:
    waarde = schoon(waarde)

    if waarde == "":
        return None

    if not re.fullmatch(r"\d+", waarde):
        raise ValueError(f"Ongeldige score {waarde!r}.")

    return int(waarde)


def afgeleide_winnaar(wedstrijd: dict[str, str]) -> str:
    thuis = schoon(wedstrijd.get("thuis", ""))
    uit = schoon(wedstrijd.get("uit", ""))
    thuis_score = score_getal(wedstrijd.get("voorspeld_thuis", ""))
    uit_score = score_getal(wedstrijd.get("voorspeld_uit", ""))
    ingevulde_winnaar = schoon(wedstrijd.get("winnaar", ""))

    if thuis_score is not None and uit_score is not None:
        if thuis_score > uit_score:
            return thuis
        if uit_score > thuis_score:
            return uit

    if ingevulde_winnaar:
        # Een placeholder is nog geen echte winnaar.
        if normaliseer(ingevulde_winnaar).startswith("winnaar "):
            return ""

        return ingevulde_winnaar

    return ""


def los_bracket_op(
    wedstrijden: list[dict[str, str]],
    bron: Path,
) -> list[dict[str, str]]:
    """Vul teams in volgende rondes automatisch vanuit vorige winnaars."""

    per_sleutel = {
        (rij["ronde"], rij["wedstrijd"]): rij
        for rij in wedstrijden
    }

    for ronde in ("qf", "sf", "f"):
        ronde_wedstrijden = sorted(
            (
                rij
                for rij in wedstrijden
                if rij["ronde"] == ronde
            ),
            key=lambda rij: str(rij["wedstrijd"]),
        )

        for rij in ronde_wedstrijden:
            sleutel = (rij["ronde"], rij["wedstrijd"])
            bronnen = BRACKET.get(sleutel)

            if bronnen is None:
                continue

            vorige_thuis = per_sleutel.get(bronnen[0])
            vorige_uit = per_sleutel.get(bronnen[1])

            if vorige_thuis is None or vorige_uit is None:
                raise ValueError(
                    f"{bron.name}: bracketbron ontbreekt voor "
                    f"{ronde} {rij['wedstrijd']}."
                )

            voorspelde_thuis = afgeleide_winnaar(vorige_thuis)
            voorspelde_uit = afgeleide_winnaar(vorige_uit)

            huidige_thuis = schoon(rij.get("thuis", ""))
            huidige_uit = schoon(rij.get("uit", ""))

            thuis_is_placeholder = (
                huidige_thuis == ""
                or normaliseer(huidige_thuis).startswith("winnaar ")
            )
            uit_is_placeholder = (
                huidige_uit == ""
                or normaliseer(huidige_uit).startswith("winnaar ")
            )

            if thuis_is_placeholder:
                if not voorspelde_thuis:
                    raise ValueError(
                        f"{bron.name}: winnaar van "
                        f"{bronnen[0][0]} {bronnen[0][1]} is nodig om "
                        f"{ronde} {rij['wedstrijd']} automatisch in te vullen."
                    )

                rij["thuis"] = voorspelde_thuis

            if uit_is_placeholder:
                if not voorspelde_uit:
                    raise ValueError(
                        f"{bron.name}: winnaar van "
                        f"{bronnen[1][0]} {bronnen[1][1]} is nodig om "
                        f"{ronde} {rij['wedstrijd']} automatisch in te vullen."
                    )

                rij["uit"] = voorspelde_uit

            # Ook in de huidige wedstrijd de winnaar automatisch afleiden
            # wanneer de score niet gelijk is.
            huidige_winnaar = afgeleide_winnaar(rij)

            if huidige_winnaar:
                rij["winnaar"] = huidige_winnaar

    return wedstrijden


def parse_rijen(
    rijen: list[list[object]],
    bron: Path,
) -> list[dict[str, str]]:
    huidige_ronde: str | None = None
    wedstrijden: list[dict[str, str]] = []

    for ruwe_rij in rijen:
        cellen = [schoon(cel) for cel in ruwe_rij]
        samengevoegd = " ".join(cellen)

        gevonden_ronde = ronde_uit_tekst(samengevoegd)

        if gevonden_ronde:
            huidige_ronde = gevonden_ronde
            continue

        if not huidige_ronde:
            continue

        # Minimaal: code | thuis | score | uit | score
        if len(cellen) < 5:
            continue

        code = schoon(cellen[0]).upper()

        if huidige_ronde == "f" and code == "":
            code = "F"

        if not geldige_wedstrijdcode(huidige_ronde, code):
            continue

        thuis = schoon(cellen[1])
        thuis_score = schoon(cellen[2])
        uit = schoon(cellen[3])
        uit_score = schoon(cellen[4])
        winnaar = schoon(cellen[5]) if len(cellen) >= 6 else ""

        # Alleen scores valideren wanneer er iets ingevuld staat.
        score_getal(thuis_score)
        score_getal(uit_score)

        wedstrijd = {
            "ronde": huidige_ronde,
            "wedstrijd": code,
            "thuis": thuis,
            "voorspeld_thuis": thuis_score,
            "uit": uit,
            "voorspeld_uit": uit_score,
            "winnaar": winnaar,
        }

        afgeleid = afgeleide_winnaar(wedstrijd)

        if afgeleid:
            wedstrijd["winnaar"] = afgeleid

        wedstrijden.append(wedstrijd)

    # Dubbele regels verwijderen: laatste versie wint.
    uniek: dict[tuple[str, str], dict[str, str]] = {}

    for wedstrijd in wedstrijden:
        uniek[
            (wedstrijd["ronde"], wedstrijd["wedstrijd"])
        ] = wedstrijd

    wedstrijden = sorted(
        uniek.values(),
        key=lambda rij: (
            ROUND_ORDER[rij["ronde"]],
            str(rij["wedstrijd"]),
        ),
    )

    verwacht = {
        "r16": set("ABCDEFGH"),
        "qf": set("1234"),
        "sf": {"X", "Y"},
        "f": {"F"},
    }

    for ronde, codes in verwacht.items():
        gevonden = {
            rij["wedstrijd"]
            for rij in wedstrijden
            if rij["ronde"] == ronde
        }
        ontbrekend = sorted(codes - gevonden)

        if ontbrekend:
            raise ValueError(
                f"{bron.name}: ontbrekende wedstrijden in {ronde}: "
                + ", ".join(ontbrekend)
            )

    return los_bracket_op(wedstrijden, bron)


def vind_naam_docx(document: Document) -> str:
    for tabel in document.tables:
        for rij in tabel.rows:
            cellen = [schoon(cel.text) for cel in rij.cells]

            if (
                len(cellen) >= 2
                and normaliseer(cellen[0]).startswith("naam")
            ):
                return cellen[1]

    for alinea in document.paragraphs:
        match = re.search(
            r"\bnaam\s*:\s*(.+)$",
            alinea.text,
            flags=re.I,
        )

        if match:
            return match.group(1).strip()

    return ""


def lees_docx(
    bron: Path,
) -> tuple[str, list[dict[str, str]]]:
    document = Document(bron)
    naam = vind_naam_docx(document)
    rijen: list[list[object]] = []

    for tabel in document.tables:
        for rij in tabel.rows:
            rijen.append([cel.text for cel in rij.cells])

    return naam, parse_rijen(rijen, bron)


def lees_pdf(
    bron: Path,
) -> tuple[str, list[dict[str, str]]]:
    naam = ""
    rijen: list[list[object]] = []

    with pdfplumber.open(bron) as pdf:
        volledige_tekst = "\n".join(
            pagina.extract_text() or ""
            for pagina in pdf.pages
        )

        match = re.search(
            r"\bnaam\s*:\s*([^\n\r]+)",
            volledige_tekst,
            flags=re.I,
        )

        if match:
            naam = match.group(1).strip()

        for pagina in pdf.pages:
            for tabel in pagina.extract_tables() or []:
                rijen.extend(tabel or [])

    return naam, parse_rijen(rijen, bron)


def lees_formulier(
    bron: Path,
) -> tuple[str, list[dict[str, str]]]:
    if bron.suffix.casefold() == ".docx":
        return lees_docx(bron)

    if bron.suffix.casefold() == ".pdf":
        return lees_pdf(bron)

    raise ValueError(
        f"Bestandstype {bron.suffix!r} wordt niet ondersteund."
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--data-dir",
        type=Path,
        default=Path("data/fase3"),
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    data_dir: Path = args.data_dir
    formulieren_dir = data_dir / "formulieren"
    punten_pad = data_dir / "punten_tot_fase2.csv"
    uitvoer_pad = data_dir / "voorspellingen_fase3.csv"

    try:
        bekende_namen = {
            normaliseer(rij.get("naam", "")): schoon(rij.get("naam", ""))
            for rij in lees_csv(punten_pad)
            if schoon(rij.get("naam", ""))
        }

        if not bekende_namen:
            raise ValueError(
                "punten_tot_fase2.csv bevat geen deelnemers."
            )

        bestaande_rijen = lees_csv(uitvoer_pad)
        per_deelnemer: dict[
            str,
            list[dict[str, str]],
        ] = {}

        for rij in bestaande_rijen:
            naam = schoon(rij.get("naam", ""))

            if naam:
                per_deelnemer.setdefault(
                    normaliseer(naam),
                    [],
                ).append(dict(rij))

        bestanden = sorted(
            (
                pad
                for pad in formulieren_dir.iterdir()
                if pad.is_file()
                and pad.suffix.casefold() in {".docx", ".pdf"}
            ),
            key=lambda pad: (
                pad.stat().st_mtime,
                pad.name.casefold(),
            ),
        )

        fouten = 0

        for bron in bestanden:
            try:
                formuliernaam, wedstrijden = lees_formulier(bron)

                if not formuliernaam:
                    raise ValueError(
                        f"{bron.name}: geen naam gevonden."
                    )

                officiele_naam = koppel_naam(
                    formuliernaam,
                    bekende_namen,
                )

                nieuwe_rijen = [
                    {
                        "naam": officiele_naam,
                        **wedstrijd,
                        "bronbestand": bron.name,
                    }
                    for wedstrijd in wedstrijden
                ]

                # Laatste geldige formulier van een deelnemer wint.
                per_deelnemer[
                    normaliseer(officiele_naam)
                ] = nieuwe_rijen

                print(
                    f"OK: {bron.name} -> {officiele_naam} "
                    f"({len(nieuwe_rijen)} wedstrijden)"
                )

            except Exception as fout:
                fouten += 1
                print(f"FOUT: {fout}", file=sys.stderr)

        alle_rijen = [
            rij
            for rijen in per_deelnemer.values()
            for rij in rijen
        ]

        alle_rijen.sort(
            key=lambda rij: (
                normaliseer(rij["naam"]),
                ROUND_ORDER[rij["ronde"]],
                str(rij["wedstrijd"]),
            )
        )

        uitvoer_pad.parent.mkdir(parents=True, exist_ok=True)

        with uitvoer_pad.open(
            "w",
            encoding="utf-8-sig",
            newline="",
        ) as bestand:
            schrijver = csv.DictWriter(
                bestand,
                fieldnames=FIELDS,
            )
            schrijver.writeheader()
            schrijver.writerows(alle_rijen)

        print(f"\nGeschreven: {uitvoer_pad.resolve()}")
        print(
            "Deelnemers in CSV: "
            f"{len(per_deelnemer)}"
        )

        if fouten:
            print(
                f"Let op: {fouten} formulier(en) bevatten een fout. "
                "De geldige formulieren zijn wel opgeslagen.",
                file=sys.stderr,
            )

        return 1 if fouten else 0

    except Exception as fout:
        print(f"FOUT: {fout}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
