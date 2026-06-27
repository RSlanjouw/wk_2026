#!/usr/bin/env python3
"""Lees WK-pouleformulieren uit en maak één compacte CSV.

Standaard:
    invoer:  data/formulieren/
    uitvoer: data/voorspellingen.csv
    fouten:  data/import_fouten.csv
    log:     data/import_log.csv

De vier slechtste nummers drie worden opgeslagen in:
    slechtste_3_1, slechtste_3_2, slechtste_3_3, slechtste_3_4

Een afwijking bij deze vier landen blokkeert de rest van het formulier niet.
De afwijking komt wel duidelijk in import_log.csv te staan.
"""

from __future__ import annotations

import argparse
import csv
import logging
import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
DEFAULT_INPUT = ROOT / "data" / "formulieren"
DEFAULT_OUTPUT = ROOT / "data" / "voorspellingen.csv"
DEFAULT_ERRORS = ROOT / "data" / "import_fouten.csv"
DEFAULT_LOG = ROOT / "data" / "import_log.csv"
DEFAULT_MAPPING = ROOT / "data" / "kolommen.csv"

# Verberg niet-fatale technische pypdf-meldingen.
logging.getLogger("pypdf").setLevel(logging.ERROR)

POOLS: dict[str, list[str]] = {
    "A": ["Zuid-Afrika", "Zuid-Korea", "Mexico", "Tsjechië"],
    "B": ["Qatar", "Canada", "Zwitserland", "Bosnië en Herzegovina"],
    "C": ["Marokko", "Brazilië", "Schotland", "Haïti"],
    "D": ["Verenigde Staten", "Paraguay", "Australië", "Turkije"],
    "E": ["Ivoorkust", "Ecuador", "Duitsland", "Curaçao"],
    "F": ["Tunesië", "Japan", "Nederland", "Zweden"],
    "G": ["Egypte", "Nieuw-Zeeland", "Iran", "België"],
    "H": ["Uruguay", "Saoedi-Arabië", "Spanje", "Kaapverdië"],
    "I": ["Senegal", "Noorwegen", "Frankrijk", "Irak"],
    "J": ["Jordanië", "Algerije", "Oostenrijk", "Argentinië"],
    "K": ["Portugal", "Colombia", "DR Congo", "Oezbekistan"],
    "L": ["Ghana", "Kroatië", "Engeland", "Panama"],
}

COLUMNS = [f"{pool}{index}" for pool in POOLS for index in range(1, 5)]
THIRD_PLACE_COLUMNS = [f"slechtste_3_{index}" for index in range(1, 5)]

TEAM_BY_COLUMN = {
    f"{pool}{index}": team
    for pool, teams in POOLS.items()
    for index, team in enumerate(teams, start=1)
}
ALL_TEAMS = set(TEAM_BY_COLUMN.values())


@dataclass
class LogEntry:
    niveau: str
    bestand: str
    deelnemer: str
    onderdeel: str
    melding: str
    waarde: str = ""


def normalize(value: str) -> str:
    value = unicodedata.normalize("NFKD", value or "")
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = value.casefold().replace("–", "-").replace("—", "-")
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return " ".join(value.split())


ALIASES: dict[str, str] = {
    normalize(team): team
    for teams in POOLS.values()
    for team in teams
}
ALIASES.update({
    "brazilie": "Brazilië",
    "belgie": "België",
    "tunesie": "Tunesië",
    "australie": "Australië",
    "curacao": "Curaçao",
    "kaapverdie": "Kaapverdië",
    "tsjechie": "Tsjechië",
    "haiti": "Haïti",
    "jordanie": "Jordanië",
    "kroatie": "Kroatië",
    "saoedi arabie": "Saoedi-Arabië",
    "saoudi arabie": "Saoedi-Arabië",
    "saudi arabie": "Saoedi-Arabië",
    "saudie arabie": "Saoedi-Arabië",
    "saudi arabia": "Saoedi-Arabië",
    "bosnie en herzegovina": "Bosnië en Herzegovina",
    "bosnie": "Bosnië en Herzegovina",
    "congo": "DR Congo",
    "d r congo": "DR Congo",
    "drc": "DR Congo",
    "algarije": "Algerije",
    "paraquay": "Paraguay",
    "verenigde staten van amerika": "Verenigde Staten",
    "usa": "Verenigde Staten",
    "zuid korea": "Zuid-Korea",
    "zuid afrika": "Zuid-Afrika",
    "nieuw zeeland": "Nieuw-Zeeland",
})


def resolve_team(value: str) -> tuple[str | None, bool]:
    """Geef (canonieke landnaam, afwijkende schrijfwijze gebruikt)."""
    key = normalize(value)
    if not key:
        return None, False

    if key in ALIASES:
        canonical = ALIASES[key]
        return canonical, key != normalize(canonical)

    # Alleen voor cellen/regels met wat extra tekst eromheen.
    candidates = [
        (len(alias), alias, team)
        for alias, team in ALIASES.items()
        if alias and re.search(rf"\b{re.escape(alias)}\b", key)
    ]
    if candidates:
        _, alias, canonical = max(candidates)
        return canonical, True

    return None, False


def canonical_team(value: str) -> str:
    team, _ = resolve_team(value)
    return team or value.strip()

def extract_all_teams_from_text(value: str) -> tuple[list[str], list[str]]:
    """Vind alle bekende landen in tekst, in leesvolgorde."""
    normalized = normalize(value)
    if not normalized:
        return [], []

    matches: list[tuple[int, int, int, str, str]] = []
    for alias, canonical in ALIASES.items():
        for match in re.finditer(rf"\b{re.escape(alias)}\b", normalized):
            matches.append(
                (match.start(), match.end(), -len(alias), alias, canonical)
            )

    matches.sort(key=lambda item: (item[0], item[2], item[1]))
    selected: list[tuple[int, int, str, str]] = []
    occupied: list[tuple[int, int]] = []

    for start, end, _, alias, canonical in matches:
        if any(
            not (end <= used_start or start >= used_end)
            for used_start, used_end in occupied
        ):
            continue
        selected.append((start, end, alias, canonical))
        occupied.append((start, end))

    selected.sort(key=lambda item: item[0])

    teams: list[str] = []
    normalizations: list[str] = []
    for _, _, alias, canonical in selected:
        if canonical not in teams:
            teams.append(canonical)
        if alias != normalize(canonical):
            normalizations.append(f"{alias} → {canonical}")

    return teams, normalizations



def validate_prediction(name: str, positions: dict[str, int]) -> None:
    missing = [team for team in TEAM_BY_COLUMN.values() if team not in positions]
    if missing:
        raise ValueError(
            f"{len(missing)} landen ontbreken: {', '.join(missing[:5])}"
        )

    for pool, teams in POOLS.items():
        values = [positions[team] for team in teams]
        if sorted(values) != [1, 2, 3, 4]:
            raise ValueError(
                f"Poule {pool} bevat geen geldige unieke posities "
                f"1,2,3,4: {values}"
            )

    if not name.strip():
        raise ValueError("naam ontbreekt")


def predicted_third_place_teams(positions: dict[str, int]) -> set[str]:
    return {team for team, position in positions.items() if position == 3}


def extract_name_docx(document: Document) -> str:
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            for index, cell in enumerate(cells):
                if normalize(cell) == "naam" and index + 1 < len(cells):
                    name = cells[index + 1].strip()
                    if name:
                        return name

    # Fallback voor het vaste oude formaat.
    if document.tables and len(document.tables[0].rows) > 0:
        cells = document.tables[0].rows[0].cells
        if len(cells) >= 2:
            return cells[1].text.strip()

    raise ValueError("naam niet gevonden")


def extract_positions_docx(document: Document) -> dict[str, int]:
    positions: dict[str, int] = {}

    # Scan alle tabellen, zodat ook een formulier met meerdere pouletabellen werkt.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]

            for team_col in range(len(cells) - 1):
                team, _ = resolve_team(cells[team_col])
                if team not in ALL_TEAMS:
                    continue

                match = re.fullmatch(r"\s*([1-4])\s*", cells[team_col + 1])
                if match:
                    positions[team] = int(match.group(1))

    return positions


def table_plain_values(table) -> list[str]:
    values: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            value = " ".join(cell.text.replace("\n", " ").split()).strip()
            if value:
                values.append(value)
    return values


def find_worst_thirds_table_docx(document: Document):
    """Vind de meest waarschijnlijke tabel met de vier slechtste nummers drie."""
    candidates: list[tuple[int, object]] = []

    for table in document.tables:
        values = table_plain_values(table)
        if not values:
            continue

        normalized_values = [normalize(value) for value in values]

        # Bonus- en pouletabellen uitsluiten.
        if any(value in {"vraag", "antwoord"} for value in normalized_values):
            continue
        if any(value.startswith("poule ") for value in normalized_values):
            continue

        recognized = 0
        for value in values:
            if re.fullmatch(r"\d+", normalize(value)):
                continue
            team, _ = resolve_team(value)
            if team:
                recognized += 1

        # Een normale tabel heeft 4 landen, soms plus nummers 1 t/m 4.
        if 2 <= recognized <= 8 and len(values) <= 12:
            score = recognized * 10
            score -= abs(recognized - 4) * 3
            score -= abs(len(table.rows) - 4)
            candidates.append((score, table))

    if not candidates:
        return None

    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def extract_worst_thirds_docx(
    document: Document,
    path: Path,
    name: str,
    positions: dict[str, int],
    log_entries: list[LogEntry],
) -> list[str]:
    table = find_worst_thirds_table_docx(document)

    raw_values: list[str] = []
    teams: list[str] = []
    unknown_values: list[str] = []
    normalized_aliases: list[str] = []

    if table is not None:
        raw_values = table_plain_values(table)

        for value in raw_values:
            if re.fullmatch(r"\s*\d+\s*", value):
                continue

            found, normalizations = extract_all_teams_from_text(value)
            if found:
                teams.extend(found)
                normalized_aliases.extend(normalizations)
            else:
                unknown_values.append(value)

    # Fallback voor formulieren waarin de vier landen niet in een losse tabel staan.
    if len(set(teams)) < 4:
        full_text_parts: list[str] = []

        for paragraph in document.paragraphs:
            value = " ".join(paragraph.text.split()).strip()
            if value:
                full_text_parts.append(value)

        for doc_table in document.tables:
            full_text_parts.extend(table_plain_values(doc_table))

        full_text = "\n".join(full_text_parts)
        normalized_full_text = normalize(full_text)

        start = re.search(
            r"welke\s+vier\s+slechtste\s+nummers?\s+drie.*?verlaten\s*:?",
            normalized_full_text,
            flags=re.I | re.S,
        )
        end = re.search(r"bonusvragen", normalized_full_text, flags=re.I)

        if start:
            section = normalized_full_text[
                start.end(): end.start() if end else len(normalized_full_text)
            ]
            found, normalizations = extract_all_teams_from_text(section)
            for team in found:
                if team not in teams:
                    teams.append(team)
            normalized_aliases.extend(normalizations)

    if not teams:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            "Geen herkenbare invoer met slechtste nummers drie gevonden.",
        ))
        return []

    return validate_and_log_worst_thirds(
        teams=teams,
        raw_values=raw_values,
        unknown_values=unknown_values,
        normalized_aliases=normalized_aliases,
        path=path,
        name=name,
        positions=positions,
        log_entries=log_entries,
    )


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_name_pdf(text: str) -> str:
    compact = " ".join(text.split())
    match = re.search(
        r"Naam:\s*(.*?)\s+Voor de poulefase",
        compact,
        flags=re.I,
    )
    if not match:
        raise ValueError("naam niet gevonden")
    return match.group(1).strip()


def extract_positions_pdf(text: str) -> dict[str, int]:
    compact = " ".join(text.split())
    start = re.search(r"Voor de poulefase.*?poule in:\s*", compact, flags=re.I)
    end = re.search(r"Tevens willen", compact, flags=re.I)

    if not start:
        raise ValueError("begin van poule-overzicht niet gevonden")

    pool_text = compact[start.end(): end.start() if end else len(compact)]
    normalized = normalize(pool_text)
    positions: dict[str, int] = {}

    for team in TEAM_BY_COLUMN.values():
        aliases = sorted(
            {
                alias
                for alias, canonical in ALIASES.items()
                if canonical == team
            },
            key=len,
            reverse=True,
        )
        for alias in aliases:
            match = re.search(
                rf"\b{re.escape(alias)}\b\s+([1-4])\b",
                normalized,
            )
            if match:
                positions[team] = int(match.group(1))
                break

    return positions


def extract_worst_thirds_pdf(
    text: str,
    path: Path,
    name: str,
    positions: dict[str, int],
    log_entries: list[LogEntry],
) -> list[str]:
    start = re.search(
        r"Tevens willen.*?verlaten\s*:",
        text,
        flags=re.I | re.S,
    )
    end = re.search(r"Bonusvragen\s*:", text, flags=re.I)

    if not start:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            "Begintekst voor slechtste nummers drie niet gevonden.",
        ))
        return []

    section = text[start.end(): end.start() if end else len(text)]
    raw_values = [
        " ".join(line.split()).strip(" :-\t")
        for line in section.splitlines()
        if " ".join(line.split()).strip(" :-\t")
    ]

    teams, normalized_aliases = extract_all_teams_from_text(section)

    unknown_values: list[str] = []
    for value in raw_values:
        found, _ = extract_all_teams_from_text(value)
        if not found and not re.fullmatch(r"\d+", normalize(value)):
            unknown_values.append(value)

    return validate_and_log_worst_thirds(
        teams=teams,
        raw_values=raw_values,
        unknown_values=unknown_values,
        normalized_aliases=normalized_aliases,
        path=path,
        name=name,
        positions=positions,
        log_entries=log_entries,
    )


def validate_and_log_worst_thirds(
    teams: list[str],
    raw_values: list[str],
    unknown_values: list[str],
    normalized_aliases: list[str],
    path: Path,
    name: str,
    positions: dict[str, int],
    log_entries: list[LogEntry],
) -> list[str]:
    # Uniek houden, maar de oorspronkelijke volgorde bewaren.
    unique_teams: list[str] = []
    duplicates: list[str] = []

    for team in teams:
        if team in unique_teams:
            duplicates.append(team)
        else:
            unique_teams.append(team)

    if len(unique_teams) != 4:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            f"Er zijn {len(unique_teams)} unieke landen gelezen; verwacht 4.",
            " | ".join(unique_teams) or "(niets gelezen)",
        ))

    if duplicates:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            "Dubbele landen gevonden.",
            " | ".join(duplicates),
        ))

    if unknown_values:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            "Niet-herkende tekst aangetroffen in het invoervak.",
            " | ".join(unknown_values),
        ))

    if normalized_aliases:
        log_entries.append(LogEntry(
            "INFO",
            path.name,
            name,
            "slechtste nummers drie",
            "Afwijkende landnamen automatisch genormaliseerd.",
            " | ".join(dict.fromkeys(normalized_aliases)),
        ))

    predicted_thirds = predicted_third_place_teams(positions)
    not_position_three = [
        team for team in unique_teams
        if team not in predicted_thirds
    ]
    if not_position_three:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            "Een of meer gekozen landen staan in de eigen poulevoorspelling "
            "niet op positie 3.",
            " | ".join(not_position_three),
        ))

    # CSV heeft altijd vier vaste kolommen. Bij minder dan vier blijven vakken leeg;
    # bij meer dan vier worden alleen de eerste vier geschreven en staat dit in de log.
    if len(unique_teams) > 4:
        log_entries.append(LogEntry(
            "WAARSCHUWING",
            path.name,
            name,
            "slechtste nummers drie",
            "Meer dan vier landen gevonden; alleen de eerste vier zijn opgeslagen.",
            " | ".join(unique_teams),
        ))

    return unique_teams[:4]


def parse_docx(
    path: Path,
    log_entries: list[LogEntry],
) -> tuple[str, dict[str, int], list[str]]:
    document = Document(path)
    if not document.tables:
        raise ValueError("onverwacht DOCX-formaat: geen tabellen gevonden")

    name = extract_name_docx(document)
    positions = extract_positions_docx(document)
    validate_prediction(name, positions)

    worst_thirds = extract_worst_thirds_docx(
        document=document,
        path=path,
        name=name,
        positions=positions,
        log_entries=log_entries,
    )
    return name, positions, worst_thirds


def parse_pdf(
    path: Path,
    log_entries: list[LogEntry],
) -> tuple[str, dict[str, int], list[str]]:
    text = extract_pdf_text(path)
    name = extract_name_pdf(text)
    positions = extract_positions_pdf(text)
    validate_prediction(name, positions)

    worst_thirds = extract_worst_thirds_pdf(
        text=text,
        path=path,
        name=name,
        positions=positions,
        log_entries=log_entries,
    )
    return name, positions, worst_thirds


def discover_files(folder: Path) -> list[Path]:
    supported = {".docx", ".pdf"}
    return sorted(
        path
        for path in folder.rglob("*")
        if path.is_file()
        and path.suffix.casefold() in supported
        and not path.name.startswith("~$")
    )


def write_mapping(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["kolom", "poule", "land"])
        for column in COLUMNS:
            writer.writerow([column, column[0], TEAM_BY_COLUMN[column]])


def write_predictions(
    path: Path,
    rows: list[dict[str, str | int]],
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = ["naam", *COLUMNS, *THIRD_PLACE_COLUMNS]
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def write_errors(path: Path, errors: list[tuple[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["bestand", "fout"])
        writer.writerows(errors)


def write_log(path: Path, entries: list[LogEntry]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow([
            "niveau",
            "bestand",
            "deelnemer",
            "onderdeel",
            "melding",
            "waarde",
        ])
        for entry in entries:
            writer.writerow([
                entry.niveau,
                entry.bestand,
                entry.deelnemer,
                entry.onderdeel,
                entry.melding,
                entry.waarde,
            ])


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--invoer",
        type=Path,
        default=DEFAULT_INPUT,
        help="map met DOCX- en PDF-formulieren",
    )
    parser.add_argument(
        "--uitvoer",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="doelbestand voor de compacte CSV",
    )
    args = parser.parse_args(argv)

    args.invoer.mkdir(parents=True, exist_ok=True)
    files = discover_files(args.invoer)

    if not files:
        print(f"Geen DOCX- of PDF-formulieren gevonden in: {args.invoer}")
        print("Plaats de formulieren in die map en voer het script opnieuw uit.")
        write_mapping(DEFAULT_MAPPING)
        return 1

    rows: list[dict[str, str | int]] = []
    errors: list[tuple[str, str]] = []
    log_entries: list[LogEntry] = []
    names_seen: set[str] = set()

    for path in files:
        log_start = len(log_entries)

        try:
            if path.suffix.casefold() == ".docx":
                name, positions, worst_thirds = parse_docx(path, log_entries)
            else:
                name, positions, worst_thirds = parse_pdf(path, log_entries)

            name_key = normalize(name)
            if name_key in names_seen:
                raise ValueError(f"dubbele deelnemer: {name}")
            names_seen.add(name_key)

            row: dict[str, str | int] = {"naam": name}

            for column in COLUMNS:
                row[column] = positions[TEAM_BY_COLUMN[column]]

            for index, column in enumerate(THIRD_PLACE_COLUMNS):
                row[column] = (
                    worst_thirds[index]
                    if index < len(worst_thirds)
                    else ""
                )

            rows.append(row)

            new_entries = log_entries[log_start:]
            warning_count = sum(
                entry.niveau == "WAARSCHUWING"
                for entry in new_entries
            )

            if warning_count:
                print(
                    f"OK  {path.name} -> {name} "
                    f"({warning_count} waarschuwing(en))"
                )
            else:
                print(f"OK  {path.name} -> {name}")

        except Exception as exc:
            errors.append((path.name, str(exc)))
            print(f"FOUT {path.name}: {exc}", file=sys.stderr)

    rows.sort(key=lambda row: normalize(str(row["naam"])))
    write_predictions(args.uitvoer, rows)
    write_mapping(DEFAULT_MAPPING)
    write_errors(DEFAULT_ERRORS, errors)
    write_log(DEFAULT_LOG, log_entries)

    warnings = sum(
        entry.niveau == "WAARSCHUWING"
        for entry in log_entries
    )

    print()
    print(f"Klaar: {len(rows)} deelnemers geschreven naar {args.uitvoer}")
    print(f"Importlog: {DEFAULT_LOG}")

    if warnings:
        print(
            f"LET OP: {warnings} waarschuwing(en) gevonden. "
            f"Controleer {DEFAULT_LOG.name}."
        )

    if errors:
        print(
            f"Controleer {DEFAULT_ERRORS.name}: "
            f"{len(errors)} bestand(en) konden niet worden verwerkt."
        )

    return 0 if rows else 1


if __name__ == "__main__":
    raise SystemExit(main())
